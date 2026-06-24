import re
import sys
from pathlib import Path

sys.path.insert(0, "E:/Desktop/codex/output/python_deps")

import markdown
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


OUT = Path("E:/Desktop/codex/output")
MD_PATH = OUT / "iqcot_multiphase_iek_paper_v3_simulink_validated.md"
HTML_PATH = OUT / "iqcot_multiphase_iek_paper_v3_simulink_validated.html"
DOCX_PATH = OUT / "iqcot_multiphase_iek_article_v3_simulink_validated.docx"


def markdown_to_html(md: str, path: Path) -> None:
    body = markdown.markdown(md, extensions=["tables", "fenced_code"])
    html = f"""<!doctype html>
<html lang="zh-CN">
<head>
<meta charset="utf-8">
<title>四相数字 IQCOT IEK 论文 v3</title>
<style>
body {{ font-family: "Microsoft YaHei", "Noto Sans CJK SC", Arial, sans-serif; line-height: 1.65; max-width: 980px; margin: 40px auto; padding: 0 24px; color: #111; }}
pre, code {{ font-family: Consolas, "Courier New", monospace; }}
pre {{ background: #f6f8fa; padding: 12px; overflow-x: auto; }}
table {{ border-collapse: collapse; width: 100%; margin: 16px 0; }}
th, td {{ border: 1px solid #ccc; padding: 6px 8px; vertical-align: top; }}
img {{ max-width: 100%; display: block; margin: 16px auto; }}
h1, h2, h3 {{ line-height: 1.3; }}
</style>
</head>
<body>
{body}
</body>
</html>
"""
    path.write_text(html, encoding="utf-8")


def add_markdown_table(doc: Document, lines: list[str], start: int) -> int:
    rows = []
    i = start
    while i < len(lines) and lines[i].strip().startswith("|"):
        cells = [c.strip() for c in lines[i].strip().strip("|").split("|")]
        if not all(set(c) <= set("-: ") for c in cells):
            rows.append(cells)
        i += 1

    if rows:
        table = doc.add_table(rows=len(rows), cols=max(len(r) for r in rows))
        table.style = "Table Grid"
        for r_idx, row in enumerate(rows):
            for c_idx, cell in enumerate(row):
                table.cell(r_idx, c_idx).text = cell
        doc.add_paragraph()
    return i


def markdown_to_docx(md: str, path: Path) -> None:
    doc = Document()
    styles = doc.styles
    styles["Normal"].font.name = "Microsoft YaHei"
    styles["Normal"].font.size = Pt(10.5)

    lines = md.splitlines()
    in_code = False
    code_buf: list[str] = []
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if stripped.startswith("```"):
            if in_code:
                p = doc.add_paragraph()
                run = p.add_run("\n".join(code_buf))
                run.font.name = "Consolas"
                run.font.size = Pt(9)
                code_buf = []
                in_code = False
            else:
                in_code = True
            i += 1
            continue

        if in_code:
            code_buf.append(line)
            i += 1
            continue

        if not stripped:
            i += 1
            continue

        if stripped.startswith("|"):
            i = add_markdown_table(doc, lines, i)
            continue

        img = re.match(r"!\[(.*?)\]\((.*?)\)", stripped)
        if img:
            alt, img_path = img.group(1), img.group(2)
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            try:
                run = p.add_run()
                run.add_picture(img_path, width=Inches(5.9))
                cap = doc.add_paragraph(alt)
                cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            except Exception:
                doc.add_paragraph(f"[图像未能嵌入: {img_path}]")
            i += 1
            continue

        if stripped.startswith("#"):
            level = len(stripped) - len(stripped.lstrip("#"))
            text = stripped[level:].strip()
            doc.add_heading(text, level=min(level, 4))
        elif stripped.startswith("- "):
            doc.add_paragraph(stripped[2:], style="List Bullet")
        elif re.match(r"^\d+\.\s+", stripped):
            doc.add_paragraph(re.sub(r"^\d+\.\s+", "", stripped), style="List Number")
        else:
            doc.add_paragraph(stripped)
        i += 1

    doc.save(path)


def main() -> None:
    md = MD_PATH.read_text(encoding="utf-8")
    markdown_to_html(md, HTML_PATH)
    markdown_to_docx(md, DOCX_PATH)
    print(f"MD={MD_PATH}")
    print(f"HTML={HTML_PATH}")
    print(f"DOCX={DOCX_PATH}")


if __name__ == "__main__":
    main()
