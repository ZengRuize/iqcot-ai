import csv
import re
import statistics
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

try:
    import markdown
except ModuleNotFoundError:
    sys.path.append("E:/Desktop/codex/output/python_deps")
    import markdown


OUT = Path("E:/Desktop/codex/output")
V6_MD = OUT / "iqcot_multiphase_iek_paper_v6_pis_iek_comprehensive.md"
MD_PATH = OUT / "iqcot_multiphase_iek_paper_v7_validated_budgeted.md"
HTML_PATH = OUT / "iqcot_multiphase_iek_paper_v7_validated_budgeted.html"
DOCX_PATH = OUT / "iqcot_multiphase_iek_article_v7_validated_budgeted.docx"


def read_csv(name):
    with (OUT / name).open(newline="", encoding="utf-8") as f:
        return list(csv.DictReader(f))


def fnum(x, digits=4):
    x = float(x)
    if x == 0:
        return "0"
    if abs(x) >= 1000 or abs(x) < 0.001:
        return f"{x:.{digits}e}"
    return f"{x:.{digits}f}".rstrip("0").rstrip(".")


def build_addendum():
    fd = read_csv("iqcot_simulink_perphase_fd_jacobian.csv")
    mc = read_csv("iqcot_pis_iek_monte_carlo_summary.csv")

    lambda_m2 = [r for r in fd if r["actuator"] == "Lambda" and r["pattern"] == "m2_alt"]
    ton_m2 = [r for r in fd if r["actuator"] == "Ton" and r["pattern"] == "m2_alt"]
    lambda_i_med = statistics.median(abs(float(r["G_m2_current_A"])) * 1e3 for r in lambda_m2)
    lambda_phi_med = statistics.median(abs(float(r["G_phase_spacing_std_ns"])) for r in lambda_m2)
    ton_i_med = statistics.median(abs(float(r["G_m2_current_A"])) * 1e3 for r in ton_m2)
    ton_phi_med = statistics.median(abs(float(r["G_phase_spacing_std_ns"])) for r in ton_m2)
    separation = ton_i_med / max(lambda_i_med, 1e-30)

    rep = [
        r for r in mc
        if r["area_bits"] == "12"
        and r["detect_clock_ns"] == "1.0"
        and r["ton_resolution_ps"] == "10"
        and r["comp_delay_sigma_ns"] == "0.5"
    ][0]
    worst = max(mc, key=lambda r: float(r["phase_spacing_std_ns_p95"]))
    best = min(mc, key=lambda r: float(r["phase_spacing_std_ns_p95"]))

    return f"""

## 15. v7 新增验证一：Simulink 逐相有限差分 Jacobian

根据审稿式建议，v7 不再只依赖解析 PIS-IEK 脚本，而是在逐相面积触发 Simulink 副本上构建了 trim 版模型：

```text
E:/Desktop/codex/output/simulink_iek/four_phase_iek_perphase_trim.slx
```

该副本仍然不修改用户原始 `four_phase.slx`，只在已有逐相面积副本上暴露 `Lambda1..Lambda4` 与 `Ton_trim1..Ton_trim4`。模型 XML 抽查确认，四个面积阈值常数已经分别绑定到 `Lambda1..Lambda4`，四个 `Ton_trim` 常数已经接入 `IQCOT_Ton_Adapter/Ton_Sum1..4`。因此，本节有限差分直接作用于 Simulink 开关模型中的执行量，而不是解析脚本里的替代变量。

有限差分实验采用 3 个负载点、2 个空间图样和 2 类执行量，共 27 个原始 Simulink 样本与 12 行中心差分 Jacobian。`Lambda` 扰动幅值为 `0.10 Lambda_area = 6e-11 V*s`；`Ton` 扰动幅值采用 `4 ns`。这里不使用 20 ps 作为 Simulink 有限差分幅值，是因为该开关模型包含离散时序/事件检测分辨率，ps 级 `Ton` 修正在部分工况中会被有效时序量化吞没。该现象本身也是数字控制建模必须保留的工程约束。

| 通道 | Simulink 有限差分中值 | 相位间隔代价中值 | 解释 |
|---|---:|---:|---|
| `Lambda_m2 -> m2 current` | `{lambda_i_med:.6g} mA/(1e-13 V*s)` | `{lambda_phi_med:.6g} ns/(1e-13 V*s)` | 电流通道极弱，主要不是 DC 均流旋钮 |
| `Ton_m2 -> m2 current` | `{ton_i_med:.6g} mA/(0.1 ns)` | `{ton_phi_med:.6g} ns/(0.1 ns)` | 电流通道强，是主要均流执行量 |

按 m2 电流投影中值估算，`Ton_m2` 与 `Lambda_m2` 的电流通道强度相差约 `{separation:.3g}` 倍。这个数量级小于理想解析模型中“数个到五个数量级”的差距，原因是 Simulink 副本包含离散时序、有限仿真窗口、负载点差异和面积触发工程近似；但方向结论保持一致：`Lambda_diff` 不是强 DC current-sharing 执行量，`Ton_diff` 才是强均流通道。

![Fig. 18. Simulink per-phase finite-difference Jacobian](E:/Desktop/codex/output/figures/fig18_simulink_fd_jacobian.png)

## 16. v7 新增验证二：数字量化与检测延迟 Monte Carlo 预算

为回应“仿真数据点太少”和“数字实现预算不足”的问题，v7 增加了 PIS-IEK 事件域 Monte Carlo。该实验扫描：

| 参数 | 扫描值 |
|---|---|
| 面积阈值位宽 | `10, 12, 14, 16 bit` |
| 检测时钟 | `0.5, 1, 2, 5 ns` |
| `Ton` 分辨率 | `5, 10, 20, 50 ps` |
| 比较器随机延迟标准差 | `0, 0.5, 1, 2 ns` |
| 每个组合随机种子 | `16` |

总计生成 `4096` 行随机样本和 `256` 行聚合统计。代表性工况 `12 bit / 1 ns clock / 10 ps Ton / 0.5 ns delay sigma` 的结果为：

| 指标 | 均值 | 95 分位 |
|---|---:|---:|
| wait jitter rms | `{fnum(rep['wait_jitter_rms_ns_mean'])} ns` | `{fnum(rep['wait_jitter_rms_ns_p95'])} ns` |
| phase-spacing std | `{fnum(rep['phase_spacing_std_ns_mean'])} ns` | `{fnum(rep['phase_spacing_std_ns_p95'])} ns` |
| current-sharing rms | `{fnum(rep['current_share_rms_mA_mean'])} mA` | `{fnum(rep['current_share_rms_mA_p95'])} mA` |
| event-level Vout rms | `{fnum(rep['vout_event_rms_mV_mean'])} mV` | `{fnum(rep['vout_event_rms_mV_p95'])} mV` |

最差聚合工况为 `bits={worst['area_bits']}, clock={worst['detect_clock_ns']} ns, Ton={worst['ton_resolution_ps']} ps, delay_sigma={worst['comp_delay_sigma_ns']} ns`，其 phase-spacing std 的 95 分位为 `{fnum(worst['phase_spacing_std_ns_p95'])} ns`。最好聚合工况为 `bits={best['area_bits']}, clock={best['detect_clock_ns']} ns, Ton={best['ton_resolution_ps']} ps, delay_sigma={best['comp_delay_sigma_ns']} ns`，其 phase-spacing std 的 95 分位为 `{fnum(best['phase_spacing_std_ns_p95'])} ns`。

![Fig. 19. PIS-IEK Monte Carlo digital implementation budget](E:/Desktop/codex/output/figures/fig19_pis_iek_monte_carlo_budget.png)

## 17. v7 证据链更新与创新边界

v7 后，本文证据链可重新整理为四层：

1. 理论层：IEK 与 PIS-IEK 将 IQCOT 面积事件、`phase_idx`、积分 reset、`Lambda_i/Ton_i` 写入 event-to-event Jacobian。
2. 解析验证层：局部灵敏度、模态投影、幅值扫描和 lifted frequency response 验证 PIS-IEK 的小信号适用范围。
3. 电路交叉验证层：Simulink 逐相面积触发副本与 trim 副本给出静态负载、面积差模和 `Ton` 差模有限差分证据。
4. 数字实现预算层：Monte Carlo 将面积位宽、检测时钟、Ton 分辨率和比较器延迟映射为 wait jitter、phase-spacing 与均流误差。

创新边界仍需克制。本文不声称首次提出 IQCOT、COT sampled-data、saltation/Poincare 方法或任意 duty 的多相 COT 通用模型。本文可主张的创新是：面向四相数字 IQCOT 非重叠工作区，把面积积分事件、相索引调度、积分 reset、数字量化与受限均流执行量统一到可计算的小信号设计框架中，并用解析大样本与 Simulink 副本交叉验证 `Lambda_diff` 与 `Ton_diff` 的执行通道差异。

## 18. v7 审稿式自查

| 审稿风险 | v7 后状态 | 仍需保留的边界 |
|---|---|---|
| 验证点太少 | 已增加 4096 行 Monte Carlo、256 行聚合统计、27 个 Simulink 样本和 12 行有限差分 Jacobian | Monte Carlo 是事件域预算，不等同于硬件实测 |
| 只在解析脚本中成立 | 已新增 trim 版 Simulink 副本，直接扰动 `Lambda1..4` 与 `Ton_trim1..4` | Simulink 副本仍是工程近似，未纳入完整 PCB/驱动非理想 |
| `Ton` 小信号与数字分辨率冲突 | 已发现 ps 级 `Ton` 在 Simulink 中可能被有效时序量化吞没，改用 4 ns 有限步长验证方向 | Simulink 有限差分是有限步长方向验证，不是解析微分等式 |
| 工作点范围偏窄 | 已补 20/40/50 A 有限差分与 20--50 A 静态负载扫点 | 仍聚焦 12 V 到 1 V、四相、非 phase-overlap 区域 |
| 外环/硬件链路不足 | 已将其列入局限，数字预算提供下一步设计约束 | 仍未完成实物、FPGA/ASIC、phase-overlap 扩展 |

因此，v7 稿件更适合作为硕士论文或中文期刊论文的完整版本；若目标是 TPEL/JESTPE，还需要补充 phase-overlap 区域、硬件链路和实测/PLECS 交叉验证。

## 19. v7 数据与脚本清单

| 类型 | 文件 |
|---|---|
| Simulink trim 副本构建 | `E:/Desktop/codex/output/iqcot_build_iek_perphase_trim_model.m` |
| Simulink trim 副本 | `E:/Desktop/codex/output/simulink_iek/four_phase_iek_perphase_trim.slx` |
| Simulink 有限差分脚本 | `E:/Desktop/codex/output/iqcot_simulink_perphase_fd_validation.m` |
| Simulink 原始样本 | `E:/Desktop/codex/output/iqcot_simulink_perphase_fd_samples.csv` |
| Simulink Jacobian | `E:/Desktop/codex/output/iqcot_simulink_perphase_fd_jacobian.csv` |
| Simulink 报告 | `E:/Desktop/codex/output/iqcot_simulink_perphase_fd_validation_report.md` |
| Monte Carlo 脚本 | `E:/Desktop/codex/output/iqcot_pis_iek_monte_carlo_budget.py` |
| Monte Carlo 明细 | `E:/Desktop/codex/output/iqcot_pis_iek_monte_carlo_detail.csv` |
| Monte Carlo 聚合 | `E:/Desktop/codex/output/iqcot_pis_iek_monte_carlo_summary.csv` |
| Monte Carlo 报告 | `E:/Desktop/codex/output/iqcot_pis_iek_monte_carlo_budget_report.md` |
| 新增图表 | `E:/Desktop/codex/output/figures/fig18_simulink_fd_jacobian.png`, `E:/Desktop/codex/output/figures/fig19_pis_iek_monte_carlo_budget.png` |
"""


def markdown_to_html(md: str, path: Path) -> None:
    body = markdown.markdown(md, extensions=["tables", "fenced_code"])
    html = f"""<!doctype html>
<html lang="zh-CN">
<head>
<meta charset="utf-8">
<title>四相数字 IQCOT IEK/PIS-IEK 论文 v7</title>
<style>
body {{ font-family: "Microsoft YaHei", "Noto Sans CJK SC", Arial, sans-serif; line-height: 1.65; max-width: 980px; margin: 40px auto; padding: 0 24px; color: #111; }}
pre, code {{ font-family: Consolas, "Courier New", monospace; }}
pre {{ background: #f6f8fa; padding: 12px; overflow-x: auto; }}
table {{ border-collapse: collapse; width: 100%; margin: 16px 0; }}
th, td {{ border: 1px solid #ccc; padding: 6px 8px; vertical-align: top; }}
img {{ max-width: 100%; display: block; margin: 16px auto; }}
h1, h2, h3 {{ line-height: 1.3; }}
</style>
</head>
<body>
{body}
</body>
</html>
"""
    path.write_text(html, encoding="utf-8")


def add_markdown_table(doc: Document, lines: list[str], start: int) -> int:
    rows = []
    i = start
    while i < len(lines) and lines[i].strip().startswith("|"):
        cells = [c.strip() for c in lines[i].strip().strip("|").split("|")]
        if not all(set(c) <= set("-: ") for c in cells):
            rows.append(cells)
        i += 1
    if rows:
        table = doc.add_table(rows=len(rows), cols=max(len(r) for r in rows))
        table.style = "Table Grid"
        for r_idx, row in enumerate(rows):
            for c_idx, cell in enumerate(row):
                table.cell(r_idx, c_idx).text = cell
        doc.add_paragraph()
    return i


def markdown_to_docx(md: str, path: Path) -> None:
    doc = Document()
    styles = doc.styles
    styles["Normal"].font.name = "Microsoft YaHei"
    styles["Normal"].font.size = Pt(10.5)

    lines = md.splitlines()
    in_code = False
    code_buf: list[str] = []
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()
        if stripped.startswith("```"):
            if in_code:
                p = doc.add_paragraph()
                run = p.add_run("\n".join(code_buf))
                run.font.name = "Consolas"
                run.font.size = Pt(9)
                code_buf = []
                in_code = False
            else:
                in_code = True
            i += 1
            continue
        if in_code:
            code_buf.append(line)
            i += 1
            continue
        if not stripped:
            i += 1
            continue
        if stripped.startswith("|"):
            i = add_markdown_table(doc, lines, i)
            continue
        img = re.match(r"!\[(.*?)\]\((.*?)\)", stripped)
        if img:
            alt, img_path = img.group(1), img.group(2)
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            try:
                run = p.add_run()
                run.add_picture(img_path, width=Inches(5.9))
                cap = doc.add_paragraph(alt)
                cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            except Exception:
                doc.add_paragraph(f"[图像未能嵌入: {img_path}]")
            i += 1
            continue
        if stripped.startswith("#"):
            level = len(stripped) - len(stripped.lstrip("#"))
            doc.add_heading(stripped[level:].strip(), level=min(level, 4))
        elif stripped.startswith("- "):
            doc.add_paragraph(stripped[2:], style="List Bullet")
        elif re.match(r"^\d+\.\s+", stripped):
            doc.add_paragraph(re.sub(r"^\d+\.\s+", "", stripped), style="List Number")
        else:
            doc.add_paragraph(stripped)
        i += 1
    doc.save(path)


def main():
    md = V6_MD.read_text(encoding="utf-8").rstrip() + build_addendum()
    MD_PATH.write_text(md, encoding="utf-8")
    markdown_to_html(md, HTML_PATH)
    markdown_to_docx(md, DOCX_PATH)
    print(f"MD={MD_PATH}")
    print(f"HTML={HTML_PATH}")
    print(f"DOCX={DOCX_PATH}")


if __name__ == "__main__":
    main()
